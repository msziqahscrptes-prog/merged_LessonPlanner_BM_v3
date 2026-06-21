import streamlit as st
import google.generativeai as genai
from docx import Document
from docx.shared import Pt, Inches
from io import BytesIO

# --- 1. KONFIGURASI HALAMAN ---
st.set_page_config(page_title="Perancang Integrasi Pengajaran", layout="wide")
st.title("🎓 Integrasi Penjana Rancangan Mengajar Harian")

# --- INPUT KUNCI API (DI BAHAGIAN ATAS) ---
user_api_key = st.text_input(
    "🔑 Masukkan Kunci API Gemini Anda:", 
    type="password", 
    help="Dapatkan kunci API anda dari Google AI Studio menggunakan akaun Gmail anda."
)

# Fungsi untuk memeriksa dan memuatkan model secara dinamik
def get_working_model(api_key):
    try:
        genai.configure(api_key=api_key)
        for m in genai.list_models():
            if 'generateContent' in m.supported_generation_methods:
                return m.name
    except Exception as e:
        st.error(f"Ralat Kunci API atau Masalah Sambungan: {str(e)}")
        return None
    return "models/gemini-1.5-flash"  # Sandaran Lalai


# Memproses pengesahan model
selected_model_name = None
if user_api_key:
    selected_model_name = get_working_model(user_api_key)
    if selected_model_name:
        st.info(f"Sistem disambungkan. Model Aktif: {selected_model_name}")
else:
    st.warning("⚠️ Sila masukkan Kunci API Gemini peribadi anda di atas untuk bermula.")


# --- 2. LOGIK AI (BAHASA MELAYU PENUH & SENARAI NOMBOR) ---
def generate_advanced_plan_bm(topic, syllabus, extra_context, api_key, model_name):
    genai.configure(api_key=api_key)
    model = genai.GenerativeModel(model_name)
    
    prompt = f"""
    Topik: {topic}. Kod Silibus: {syllabus}. Konteks Tambahan: {extra_context}.
    Janakan rancangan mengajar profesional dalam Bahasa Melayu sepenuhnya.
    
    PERATURAN UTAMA:
    1. JANGAN gunakan simbol dwi-asterisk (**) sama sekali dalam teks output.
    2. JANGAN gunakan senarai peluru (bullet points). Gunakan sistem bernombor (1, 2, 3...) untuk semua senarai kriteria.
    3. JANGAN gunakan perkataan 'Murid' dalam teks, digantikan dengan perkataan 'Pelajar' dalam keseluruhan teks.
    
    Gunakan penanda (markers) HURUF BESAR yang tepat ini untuk struktur dokumen:
    
    SECTION: TOPIK PELAJARAN
    {topic}
    
    SECTION: OBJEKTIF PELAJARAN
    [4 mata huraian bernombor 1 hingga 4]
    
    SECTION: HASIL PELAJARAN
    [4 mata huraian bernombor 1 hingga 4]
    
    SECTION: KRITERIA KEJAYAAN
    [4 mata huraian bernombor 1 hingga 4]
    
    SECTION: PRASYARAT
    [1 mata huraian utama bernombor 1]
    
    SECTION: KATA KUNCI
    [6 item perkataan bernombor 1 hingga 6]
    
    SECTION: KBAT
    [4 domain utama dari Taksonomi Bloom bernombor 1 hingga 4]
    
    SECTION: KEWARGANEGARAAN DIGITAL
    [4 mata bernombor 1 hingga 4 mengenai penggunaan teknologi secara beretika/Chromebook/Canva/YouTube]

    SECTION: KANDUNGAN PEMBUKAAN PELAJARAN
    [Aktiviti set induksi dan pelan transisi]

    SECTION: STRATEGI DIFERENSIASI (HIJAU)
    1. HA (Murid Pencapaian Tinggi): [1 aktiviti mencabar dalam bentuk item bernombor]

    SECTION: STRATEGI DIFERENSIASI (KUNING)
    1. MA (Murid Pencapaian Sederhana): [1 aktiviti teras dalam bentuk item bernombor]

    SECTION: STRATEGI DIFERENSIASI (MERAH)
    1. LA (Murid Pencapaian Rendah): [1 aktiviti sokongan berperancah dalam bentuk item bernombor]

    SECTION: AKTIVITI PEMBELAJARAN TERADUN SATU (15 MINIT)
    1. Aktiviti 1: [Huraian Ringkas]
    2. Persediaan Guru: [Langkah demi langkah dalam senarai bernombor]
    3. Objektif: [3 mata huraian bernombor]
    4. Tugasan Murid: [Langkah kerja murid dalam senarai bernombor]

    SECTION: AKTIVITI PEMBELAJARAN TERADUN DUA (15 MINIT)
    1. Aktiviti 2: [Huraian Ringkas]
    2. Persediaan Guru: [Langkah demi langkah dalam senarai bernombor]
    3. Objektif: [3 mata huraian bernombor]
    4. Tugasan Murid: [Langkah kerja murid dalam senarai bernombor]

    SECTION: GRID ALIRAN PEDATI
    [Janakan kandungan untuk 4 blok pedagogi menggunakan susun atur di bawah. Ringkas dan praktikal.]
    BLOCK_START: P: PREPARATION / PENGEGETAHUAN SEDIA ADA (LEARN)
    GURU: [Langkah tindakan selaras dengan topik]
    MURID: [Tugasan murid/penggunaan Chromebook yang selaras dengan topik]
    BLOCK_END
    
    BLOCK_START: E: ENGAGE / ENGAGE (EXPLORE)
    GURU: [Langkah tindakan selaras dengan topik]
    MURID: [Tugasan murid/penggunaan Chromebook yang selaras dengan topik]
    BLOCK_END

    BLOCK_START: D.A: DELIVER AND APPLY / PERKEMBANGAN DAN APLIKASI
    GURU: [Langkah tindakan selaras dengan topik]
    MURID: [Tugasan murid/penggunaan Chromebook yang selaras dengan topik]
    BLOCK_END

    BLOCK_START: T.I: TEST AND EVALUATE / TAKSIR DAN PENAMBAHBAIKAN
    GURU: [Langkah tindakan selaras dengan topik]
    MURID: [Tugasan murid/penggunaan Chromebook yang selaras dengan topik]
    BLOCK_END
    
    SECTION: PLENARY (TIKET KELUAR)
    [Aktiviti penutup 2-3 minit]

    SECTION: KERJA RUMAH
    [Tugasan berdasarkan topik]

    SECTION: CADANGAN TUGASAN SETERUSNYA
    [Aktiviti set induksi hari esok dan perancangan transisi]
    """
    try:
        response = model.generate_content(prompt)
        if response.candidates and response.candidates[0].content.parts:
            return response.text.replace("**", "")
        else:
            return "⚠️ AI mengembalikan respons kosong. Sila tunggu 60 saat dan cuba lagi (Had Quota Minit Penuh)."
    except Exception as e:
        return f"Ralat Sistem: {str(e)}"


# --- 3. LOGIK EKSPORT WORD ---
def create_word_export_bm(topic, syllabus, text):
    doc = Document()
    doc.add_heading(f'PTES RANCANGAN MENGAJAR: {topic.upper()}', 0)

    # Jadual Admin Atas
    admin_table = doc.add_table(rows=3, cols=4)
    admin_table.style = 'Table Grid'
    labels = [["Minggu No:", "Tarikh:"], ["Bilangan Murid:", "Hari:"], ["Tempat/Makmal:", "Durasi Kelas:"]]
    for r in range(3):
        admin_table.cell(r, 0).text = labels[r][0]
        admin_table.cell(r, 2).text = labels[r][1]
    doc.add_paragraph()

    sections = text.split('SECTION:')
 
    for section in sections:
        if not section.strip(): continue
        lines = section.strip().split('\n')
        title = lines[0].strip().upper().replace("**", "")
        body_content = "\n".join(lines[1:]).strip()

        if "GRID ALIRAN PEDATI" in title or "PEDATI FLOW GRID" in title:
            doc.add_heading("PECAHAN STRUKTUR ALIRAN P.E.D.A.T.I", level=1)
            
            blocks = body_content.split("BLOCK_START:")
            for block in blocks:
                if not block.strip(): continue
                block_data = block.split("BLOCK_END")[0].strip().split('\n')
                
                heading_title = block_data[0].strip().upper().replace("**", "")
                lecturer_text = ""
                students_text = ""
                
                for line in block_data:
                    if line.upper().startswith("GURU:") or line.upper().startswith("LECTURER:"):
                        lecturer_text = line.split(":", 1)[1].strip().replace("**", "")
                    elif line.upper().startswith("MURID:") or line.upper().startswith("STUDENTS:"):
                        students_text = line.split(":", 1)[1].strip().replace("**", "")
                
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(12)
                run = p.add_run(heading_title)
                run.bold = True
                run.font.size = Pt(12)
                
                table = doc.add_table(rows=2, cols=2)
                table.style = 'Table Grid'
                
                for row in table.rows:
                    row.cells[0].width = Inches(3.25)
                    row.cells[1].width = Inches(3.25)
                
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = "Guru / Fasilitator"
                hdr_cells[1].text = "Murid / Pelajar"
                hdr_cells[0].paragraphs[0].runs[0].font.italic = True
                hdr_cells[0].paragraphs[0].runs[0].font.bold = True
                hdr_cells[1].paragraphs[0].runs[0].font.italic = True
                hdr_cells[1].paragraphs[0].runs[0].font.bold = True
                
                data_cells = table.rows[1].cells
                data_cells[0].text = lecturer_text
                data_cells[1].text = students_text
        else:
            content = body_content.replace("**", "") 
            doc.add_heading(title, level=1)
            table = doc.add_table(rows=1, cols=1)
            table.style = 'Table Grid'
            table.cell(0, 0).text = content
            doc.add_paragraph()
     
    # Pengesahan Ketua Jabatan
    doc.add_page_break()
    doc.add_heading("KELULUSAN & ULASAN KETUA JABATAN / HKP", level=1)
    hod_table = doc.add_table(rows=2, cols=2)
    hod_table.style = 'Table Grid'
    hod_table.cell(0, 0).text = "Catatan/Ulasan:"
    hod_table.rows[1].height = Pt(50)
    hod_table.cell(1, 0).text = "Tarikh:"; hod_table.cell(1, 1).text = "Tandatangan:"

    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

# --- 4. GUI UTAMA ---
st.write("---")
st.info("Sila masukkan topik pelajaran, kod silibus subjek dan maklumat tambahan seperti canva, youtube, infografik.")

c1, c2 = st.columns(2)
with c1: u_topic = st.text_input("Topik Pelajaran:")
with c2: u_syllabus = st.text_input("Kod Silibus:")
u_extra = st.text_area("Konteks Khas / Kata Kunci Tambahan (Opsional):")

if st.button("🚀 JANAKAN RANCANGAN MENGAJAR"):
    if not user_api_key:
        st.error("❌ Sila masukkan Kunci API Gemini anda di bahagian atas sebelum menjana.")
    elif not u_topic or not u_syllabus:
        st.error("❌ Sila isi maklumat bidang Topik Pelajaran dan Kod Silibus.")
    else:
        with st.spinner("AI sedang mengintegrasikan kriteria ke dalam rancangan mengajar anda..."):
            result = generate_advanced_plan_bm(u_topic, u_syllabus, u_extra, user_api_key, selected_model_name)
            st.session_state['adv_plan_out_bm'] = result

if 'adv_plan_out_bm' in st.session_state:
    st.divider()
    st.subheader("Pratonton Draf AI")
    st.text_area("Kandungan", st.session_state['adv_plan_out_bm'], height=400)
    doc_file = create_word_export_bm(u_topic, u_syllabus, st.session_state['adv_plan_out_bm'])
    st.download_button("📥 Muat Turun Versi Word (.docx)", doc_file, f"Integrasi_RPH_{u_topic}.docx")

st.markdown("---")
st.caption("Rancangan Pengajaran Harian 3.0 (BM) | Pencipta: Hjh Nurul Haziqah Hj Nordin | © 2026 Education Innovation")
